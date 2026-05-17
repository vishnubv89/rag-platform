"""
SharePoint connector via Microsoft Graph API.

Config keys:
  tenant_id       Azure AD tenant ID
  client_id       App registration client ID
  client_secret   App registration client secret
  site_url        SharePoint site URL e.g. https://acme.sharepoint.com/sites/wiki
  folder_path     (optional) library-relative folder path e.g. /Shared Documents/KB

Permissions required: Sites.Read.All (application permission)
"""
import re

import httpx

from rag_chatbot.connectors.base import BaseConnector, ConnectorDocument, RemoteDocument
from rag_chatbot.connectors.registry import register

GRAPH = "https://graph.microsoft.com/v1.0"
TOKEN_URL = "https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"


def _html_to_text(html: str) -> str:
    text = re.sub(r"<[^>]+>", " ", html)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


@register
class SharePointConnector(BaseConnector):
    connector_type = "sharepoint"

    async def _token(self) -> str:
        url = TOKEN_URL.format(tenant_id=self.config["tenant_id"])
        async with httpx.AsyncClient() as client:
            r = await client.post(url, data={
                "grant_type": "client_credentials",
                "client_id": self.config["client_id"],
                "client_secret": self.config["client_secret"],
                "scope": "https://graph.microsoft.com/.default",
            })
            r.raise_for_status()
            return r.json()["access_token"]

    async def _graph(self, token: str) -> httpx.AsyncClient:
        return httpx.AsyncClient(
            base_url=GRAPH,
            headers={"Authorization": f"Bearer {token}"},
            timeout=30,
        )

    async def _site_id(self, client: httpx.AsyncClient) -> str:
        site_url = self.config["site_url"]
        # extract hostname and site path from full URL
        m = re.match(r"https://([^/]+)(/.+)", site_url)
        if not m:
            raise ValueError(f"Invalid site_url: {site_url}")
        host, path = m.group(1), m.group(2).rstrip("/")
        r = await client.get(f"/sites/{host}:{path}")
        r.raise_for_status()
        return r.json()["id"]

    async def validate_config(self) -> tuple[bool, str]:
        required = ["tenant_id", "client_id", "client_secret", "site_url"]
        missing = [k for k in required if not self.config.get(k)]
        if missing:
            return False, f"Missing config keys: {missing}"
        try:
            token = await self._token()
            async with await self._graph(token) as client:
                await self._site_id(client)
            return True, ""
        except Exception as e:
            return False, str(e)

    async def list_documents(self) -> list[RemoteDocument]:
        token = await self._token()
        async with await self._graph(token) as client:
            site_id = await self._site_id(client)
            folder = self.config.get("folder_path", "").strip("/")
            if folder:
                r = await client.get(f"/sites/{site_id}/drive/root:/{folder}:/children")
            else:
                r = await client.get(f"/sites/{site_id}/drive/root/children")
            r.raise_for_status()
            items = r.json().get("value", [])

        docs = []
        for item in items:
            if item.get("file") and item["name"].lower().endswith((".pdf", ".docx", ".txt", ".md")):
                docs.append(RemoteDocument(
                    external_id=item["id"],
                    title=item["name"],
                    source_url=item.get("webUrl", ""),
                    updated_at=item.get("lastModifiedDateTime", ""),
                ))
        return docs

    async def fetch_document(self, external_id: str) -> ConnectorDocument:
        token = await self._token()
        async with await self._graph(token) as client:
            site_id = await self._site_id(client)
            # get item metadata
            r = await client.get(f"/sites/{site_id}/drive/items/{external_id}")
            r.raise_for_status()
            meta = r.json()
            # download content
            dl = await client.get(f"/sites/{site_id}/drive/items/{external_id}/content",
                                  follow_redirects=True)
            dl.raise_for_status()

        name = meta.get("name", "")
        raw = dl.content

        # extract text based on file type
        if name.lower().endswith(".pdf"):
            import io
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(raw))
                text = "\n".join(p.extract_text() or "" for p in reader.pages)
            except ImportError:
                text = raw.decode("latin-1", errors="replace")
        elif name.lower().endswith(".docx"):
            import io
            try:
                import docx
                doc = docx.Document(io.BytesIO(raw))
                text = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                text = raw.decode("utf-8", errors="replace")
        else:
            text = raw.decode("utf-8", errors="replace")

        return ConnectorDocument(
            external_id=external_id,
            title=name,
            text=text.strip(),
            source_url=meta.get("webUrl", ""),
            metadata={"updated_at": meta.get("lastModifiedDateTime", ""), "source": "sharepoint"},
        )
